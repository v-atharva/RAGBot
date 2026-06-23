"""Stage 2 — extract text (with page/segment spans) from each kept source file.

Routing is by extension. PDFs are page-tagged so chunks can cite a page number. Legacy ``.doc``
is converted with macOS ``textutil`` when available (no LibreOffice dependency). Image-only
files are captioned by a vision model; until an API key is configured the captioner returns a
placeholder so the rest of the pipeline can run offline.
"""

from __future__ import annotations

import subprocess
from collections.abc import Callable
from pathlib import Path

from .models import Category, ExtractedDoc, PageSpan, SourceFile

# A captioner turns an image path into a text description (vision model). None -> placeholder.
Captioner = Callable[[Path], str] | None


def extract(source: SourceFile, *, captioner: Captioner = None) -> ExtractedDoc:
    """Extract text from one classified file. ``captioner`` is an optional callable
    ``(Path) -> str`` used for diagram images; if None, a placeholder is emitted."""
    path = Path(source.path)
    ext = source.ext
    try:
        if ext == ".pdf":
            return _extract_pdf(source, path, captioner=captioner)
        if ext == ".docx":
            return _extract_docx(source, path)
        if ext == ".doc":
            return _extract_doc_legacy(source, path)
        if ext in {".txt", ".sql"}:
            return _extract_plaintext(source, path)
        if ext == ".html":
            return _extract_html(source, path)
        if ext in {".xls", ".xlsx"}:
            return _extract_spreadsheet(source, path)
        if ext in {".jpg", ".jpeg", ".png"}:
            return _extract_image(source, path, captioner=captioner)
    except Exception as exc:  # noqa: BLE001 - record, don't crash the batch
        return ExtractedDoc(source=source, text="", extractor=ext, extract_error=str(exc))

    return ExtractedDoc(
        source=source, text="", extractor="none", extract_error=f"no extractor for {ext}"
    )


def load_transcript(path: Path) -> ExtractedDoc:
    """Load a lecture transcript (``[HH:MM:SS] Speaker: text`` lines). Header lines starting
    with ``#`` are dropped. Categorized as LECTURE_TRANSCRIPT so the chunker timestamps it."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    lines = [ln for ln in raw.splitlines() if not ln.lstrip().startswith("#")]
    text = "\n".join(lines).strip()
    source = SourceFile(
        path=str(path),
        name=path.name,
        ext=path.suffix.lower(),
        size_bytes=path.stat().st_size,
        category=Category.LECTURE_TRANSCRIPT,
    )
    return ExtractedDoc(
        source=source, text=text, pages=_whole_span("transcript", text), extractor="transcript"
    )


def _whole_span(label: str, text: str) -> list[PageSpan]:
    return [PageSpan(label=label, start=0, end=len(text))]


def _extract_pdf(source: SourceFile, path: Path, *, captioner: Captioner) -> ExtractedDoc:
    import fitz  # pymupdf

    parts: list[str] = []
    pages: list[PageSpan] = []
    offset = 0
    image_only_pages = 0
    with fitz.open(path) as doc:
        for i, page in enumerate(doc, start=1):
            page_text = page.get_text("text").strip()
            if not page_text:
                image_only_pages += 1
            label = f"p.{i}"
            block = page_text + "\n\n"
            parts.append(block)
            pages.append(PageSpan(label=label, start=offset, end=offset + len(page_text)))
            offset += len(block)
        page_count = doc.page_count

    text = "".join(parts).strip()
    extractor = "pymupdf"
    error = None
    # A PDF that yielded almost no text is image-based (scanned slides / diagrams) -> caption it.
    if page_count and image_only_pages == page_count:
        caption = _caption(path, captioner)
        text = caption
        pages = _whole_span("image", text)
        extractor = "pymupdf+caption"
        error = "image-only pdf: captioned" if captioner else "image-only pdf: caption stubbed"

    return ExtractedDoc(
        source=source, text=text, pages=pages, extractor=extractor, extract_error=error
    )


def _extract_docx(source: SourceFile, path: Path) -> ExtractedDoc:
    import docx

    document = docx.Document(str(path))
    paras = [p.text for p in document.paragraphs if p.text.strip()]
    # Include simple table cell text (assignment prompts often use tables).
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    text = "\n".join(paras).strip()
    return ExtractedDoc(
        source=source, text=text, pages=_whole_span("doc", text), extractor="python-docx"
    )


def _extract_doc_legacy(source: SourceFile, path: Path) -> ExtractedDoc:
    # macOS textutil converts legacy .doc to plain text without LibreOffice.
    proc = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        capture_output=True,
        timeout=60,
    )
    if proc.returncode != 0:
        err = proc.stderr.decode("utf-8", "replace").strip()
        return ExtractedDoc(
            source=source, text="", extractor="textutil", extract_error=err or "textutil failed"
        )
    text = proc.stdout.decode("utf-8", "replace").strip()
    return ExtractedDoc(
        source=source, text=text, pages=_whole_span("doc", text), extractor="textutil"
    )


def _extract_plaintext(source: SourceFile, path: Path) -> ExtractedDoc:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    label = "sql" if source.ext == ".sql" else "txt"
    return ExtractedDoc(
        source=source, text=text, pages=_whole_span(label, text), extractor="read"
    )


def _extract_html(source: SourceFile, path: Path) -> ExtractedDoc:
    from bs4 import BeautifulSoup

    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    text = soup.get_text("\n").strip()
    return ExtractedDoc(
        source=source, text=text, pages=_whole_span("html", text), extractor="bs4"
    )


def _extract_spreadsheet(source: SourceFile, path: Path) -> ExtractedDoc:
    import pandas as pd

    sheets = pd.read_excel(path, sheet_name=None, header=None)
    parts = []
    for name, frame in sheets.items():
        parts.append(f"# sheet: {name}")
        parts.append(frame.fillna("").astype(str).to_csv(index=False, header=False))
    text = "\n".join(parts).strip()
    return ExtractedDoc(
        source=source, text=text, pages=_whole_span("sheet", text), extractor="pandas"
    )


def _extract_image(source: SourceFile, path: Path, *, captioner: Captioner) -> ExtractedDoc:
    text = _caption(path, captioner)
    return ExtractedDoc(
        source=source,
        text=text,
        pages=_whole_span("image", text),
        extractor="caption" if captioner else "caption-stub",
        extract_error=None if captioner else "caption stubbed (no captioner configured)",
    )


def _caption(path: Path, captioner: Captioner) -> str:
    if captioner is None:
        # Placeholder so stages 3-4 run offline; replaced once an API key is configured.
        return f"[IMAGE: {path.name} — caption pending vision model]"
    return captioner(path)
